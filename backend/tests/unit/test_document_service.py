"""Unit tests for document service orchestration."""

import uuid
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from app.core.exceptions import DocumentProcessingError, NotFoundError
from app.services import document_service


@pytest.fixture
def mock_db():
    """Create a mock async database session."""
    db = AsyncMock()
    db.add = MagicMock()
    db.flush = AsyncMock()
    db.delete = AsyncMock()
    return db


@pytest.fixture
def sample_docx_bytes():
    """Create a minimal Word document as bytes."""
    import io

    from docx import Document

    doc = Document()
    doc.add_paragraph("This is test content for the document service.")
    doc.add_paragraph("It contains multiple paragraphs to test chunking.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestProcessDocument:
    """Tests for the process_document function."""

    @pytest.mark.asyncio
    @patch("app.services.document_service.vector_store")
    @patch("app.services.document_service.embedding_service")
    async def test_successful_processing(
        self, mock_embed_svc, mock_vector, mock_db, sample_docx_bytes
    ):
        """Should process a document end-to-end: extract → chunk → embed → store."""
        # Mock embeddings
        mock_embed_svc.embed_texts = AsyncMock(return_value=[[0.1] * 1536])

        # Mock vector store
        mock_vector.ensure_collection = MagicMock()
        mock_vector.upsert_chunks = MagicMock(return_value=1)

        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        doc = await document_service.process_document(
            filename="test.docx",
            content_type=content_type,
            file_bytes=sample_docx_bytes,
            db=mock_db,
        )

        assert doc.status == "ready"
        assert doc.chunk_count > 0
        assert doc.filename == "test.docx"
        assert doc.file_type == "docx"
        mock_db.add.assert_called_once()
        mock_embed_svc.embed_texts.assert_awaited_once()
        mock_vector.upsert_chunks.assert_called_once()

    @pytest.mark.asyncio
    async def test_unsupported_type_rejected(self, mock_db):
        """Should raise DocumentProcessingError for unsupported MIME types."""
        with pytest.raises(DocumentProcessingError, match="Unsupported file type"):
            await document_service.process_document(
                filename="test.txt",
                content_type="text/plain",
                file_bytes=b"some text",
                db=mock_db,
            )

    @pytest.mark.asyncio
    async def test_oversized_file_rejected(self, mock_db):
        """Should raise DocumentProcessingError for files exceeding 50MB."""
        big_bytes = b"x" * (51 * 1024 * 1024)  # 51 MB
        with pytest.raises(DocumentProcessingError, match="exceeds maximum"):
            await document_service.process_document(
                filename="huge.pdf",
                content_type="application/pdf",
                file_bytes=big_bytes,
                db=mock_db,
            )

    @pytest.mark.asyncio
    async def test_empty_file_rejected(self, mock_db):
        """Should raise DocumentProcessingError for empty files."""
        with pytest.raises(DocumentProcessingError, match="empty"):
            await document_service.process_document(
                filename="empty.pdf",
                content_type="application/pdf",
                file_bytes=b"",
                db=mock_db,
            )

    @pytest.mark.asyncio
    @patch("app.services.document_service.extract_text")
    async def test_extraction_failure_sets_error_status(self, mock_extract, mock_db):
        """On extraction failure, document status should be set to 'error'."""
        mock_extract.side_effect = DocumentProcessingError("Extraction failed")

        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        with pytest.raises(DocumentProcessingError):
            await document_service.process_document(
                filename="bad.docx",
                content_type=content_type,
                file_bytes=b"corrupted data",
                db=mock_db,
            )

        # The document should have been added with error status
        added_doc = mock_db.add.call_args[0][0]
        assert added_doc.status == "error"


class TestListDocuments:
    """Tests for the list_documents function."""

    @pytest.mark.asyncio
    async def test_returns_list(self, mock_db):
        """Should return a list of documents."""
        mock_result = MagicMock()
        mock_result.scalars.return_value.all.return_value = []
        mock_db.execute = AsyncMock(return_value=mock_result)

        docs = await document_service.list_documents(mock_db)
        assert docs == []


class TestGetDocument:
    """Tests for the get_document function."""

    @pytest.mark.asyncio
    async def test_not_found_raises_error(self, mock_db):
        """Should raise NotFoundError for non-existent document."""
        mock_db.get = AsyncMock(return_value=None)
        doc_id = str(uuid.uuid4())

        with pytest.raises(NotFoundError, match="not found"):
            await document_service.get_document(doc_id, mock_db)


class TestDeleteDocument:
    """Tests for the delete_document function."""

    @pytest.mark.asyncio
    @patch("app.services.document_service.vector_store")
    async def test_delete_calls_vector_store_and_db(self, mock_vector, mock_db):
        """Should delete from both Qdrant and PostgreSQL."""
        mock_doc = MagicMock()
        mock_doc.id = uuid.uuid4()
        mock_doc.filename = "test.pdf"
        mock_db.get = AsyncMock(return_value=mock_doc)
        mock_vector.delete_by_document = MagicMock()

        doc_id = str(mock_doc.id)
        await document_service.delete_document(doc_id, mock_db)

        mock_vector.delete_by_document.assert_called_once_with(doc_id)
        mock_db.delete.assert_awaited_once_with(mock_doc)
